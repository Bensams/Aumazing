from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors

ROOT = Path(__file__).resolve().parents[1]
source = ROOT / "docs" / "Capstone_Validation_Deliverables.md"
out = ROOT / "docs" / "generated"
out.mkdir(exist_ok=True)
text = source.read_text(encoding="utf-8")

# Keep the editable form intentionally plain so advisers can print or revise it.
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(.65); sec.bottom_margin = Inches(.65)
for line in text.splitlines():
    if line.startswith("# "):
        doc.add_heading(line[2:], 0)
    elif line.startswith("## "):
        doc.add_heading(line[3:], 1)
    elif line.startswith("### "):
        doc.add_heading(line[4:], 2)
    elif line.startswith("| "):
        # Tables are represented as a readable line in the editable form; preserve all fields.
        p = doc.add_paragraph(line.replace("|", "  "))
        p.style = doc.styles["No Spacing"]
        for run in p.runs: run.font.name = "Arial"; run.font.size = Pt(8)
    elif line.strip() == "":
        doc.add_paragraph("")
    else:
        p = doc.add_paragraph(line)
        for run in p.runs: run.font.name = "Arial"; run.font.size = Pt(10)
doc.save(out / "Capstone_Validation_Forms.docx")

styles = getSampleStyleSheet()
styles.add(ParagraphStyle(name="Small", parent=styles["BodyText"], fontSize=8, leading=10))
story = []
for line in text.splitlines():
    if line.startswith("# "):
        story.append(Paragraph(line[2:], styles["Title"]))
    elif line.startswith("## "):
        story.append(Paragraph(line[3:], styles["Heading1"]))
    elif line.startswith("### "):
        story.append(Paragraph(line[4:], styles["Heading2"]))
    elif line.startswith("| "):
        cells = [c.strip().replace("`", "") for c in line.strip("|").split("|")]
        story.append(Table([[Paragraph(c, styles["Small"]) for c in cells]], colWidths=None,
                           style=TableStyle([("GRID", (0,0), (-1,-1), .25, colors.grey), ("VALIGN", (0,0), (-1,-1), "TOP")])))
    elif line.strip() == "":
        story.append(Spacer(1, 5))
    else:
        story.append(Paragraph(line.replace("`", ""), styles["BodyText"]))
pdf = out / "Capstone_Validation_Forms.pdf"
SimpleDocTemplate(str(pdf), pagesize=letter, rightMargin=.55*inch, leftMargin=.55*inch, topMargin=.5*inch, bottomMargin=.5*inch).build(story)
print(out / "Capstone_Validation_Forms.docx")
print(pdf)
